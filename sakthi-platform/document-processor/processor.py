# document_processor/processor.py
"""
Multi-format Document Processing Layer
Handles PDF, DOCX, CSV, URLs and converts to Sakthi-compatible format
"""

import io
import os
import re
import csv
import json
import requests
import pandas as pd
from typing import Dict, List, Any, Optional, Union, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging
from datetime import datetime

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    from docx import Document
    import openpyxl
    from bs4 import BeautifulSoup
    import pytesseract
    from PIL import Image
except ImportError as e:
    logging.warning(f"Some optional dependencies not installed: {e}")

# Web scraping
try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from serpapi import GoogleSearch
except ImportError as e:
    logging.warning(f"Web scraping dependencies not available: {e}")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    filename: str
    file_type: str
    size_bytes: int
    page_count: Optional[int] = None
    encoding: Optional[str] = None
    creation_date: Optional[datetime] = None
    extracted_at: datetime = datetime.now()

@dataclass
class ProcessedDocument:
    content: str
    metadata: DocumentMetadata
    structured_data: Optional[Dict[str, Any]] = None
    tables: Optional[List[Dict[str, Any]]] = None
    images: Optional[List[str]] = None  # Base64 encoded images
    confidence: float = 1.0

class DocumentProcessor:
    """Universal document processor for multiple formats"""
    
    def __init__(self, serpapi_key: Optional[str] = None):
        self.serpapi_key = serpapi_key
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.txt': self._process_text,
            '.json': self._process_json,
            '.xml': self._process_xml,
            '.html': self._process_html
        }
    
    def process_file(self, file_path: Union[str, Path, io.BytesIO], 
                    filename: Optional[str] = None) -> ProcessedDocument:
        """Process a file and extract content"""
        
        if isinstance(file_path, (str, Path)):
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            filename = filename or file_path.name
            file_extension = file_path.suffix.lower()
            file_size = file_path.stat().st_size
            
            with open(file_path, 'rb') as f:
                file_content = f.read()
        else:
            # Handle BytesIO
            file_content = file_path.read()
            file_path.seek(0)  # Reset position
            file_size = len(file_content)
            file_extension = Path(filename).suffix.lower() if filename else '.txt'
        
        # Create metadata
        metadata = DocumentMetadata(
            filename=filename or "unknown",
            file_type=file_extension,
            size_bytes=file_size
        )
        
        # Process based on file type
        processor = self.supported_formats.get(file_extension, self._process_text)
        
        try:
            if isinstance(file_path, (str, Path)):
                result = processor(file_path)
            else:
                result = processor(file_content, filename)
            
            # Update metadata
            result.metadata = metadata
            
            logger.info(f"Successfully processed {filename} ({file_extension})")
            return result
            
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            raise
    
    def process_url(self, url: str, use_serpapi: bool = False) -> ProcessedDocument:
        """Process content from URL"""
        
        if use_serpapi and self.serpapi_key:
            return self._process_url_serpapi(url)
        else:
            return self._process_url_direct(url)
    
    def process_text_input(self, text: str, context: Optional[str] = None) -> ProcessedDocument:
        """Process direct text input"""
        
        metadata = DocumentMetadata(
            filename="text_input",
            file_type="text",
            size_bytes=len(text.encode('utf-8'))
        )
        
        return ProcessedDocument(
            content=text,
            metadata=metadata,
            structured_data={"context": context} if context else None
        )
    
    def _process_pdf(self, file_path: Union[str, Path, bytes], 
                    filename: Optional[str] = None) -> ProcessedDocument:
        """Extract text and tables from PDF"""
        
        content_parts = []
        tables = []
        images = []
        
        try:
            # Use pdfplumber for better table extraction
            if isinstance(file_path, bytes):
                pdf_file = io.BytesIO(file_path)
            else:
                pdf_file = file_path
            
            with pdfplumber.open(pdf_file) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"--- Page {page_num} ---\n{text}")
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table:
                            table_dict = {
                                'page': page_num,
                                'table_index': table_idx,
                                'headers': table[0] if table else [],
                                'data': table[1:] if len(table) > 1 else []
                            }
                            tables.append(table_dict)
            
            # Fallback to PyPDF2 if pdfplumber fails
            if not content_parts:
                if isinstance(file_path, bytes):
                    pdf_file = io.BytesIO(file_path)
                else:
                    pdf_file = open(file_path, 'rb')
                
                try:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    page_count = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text:
                            content_parts.append(f"--- Page {page_num} ---\n{text}")
                finally:
                    if not isinstance(file_path, bytes):
                        pdf_file.close()
            
            content = "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            content = f"Error processing PDF: {str(e)}"
            page_count = 0
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "document.pdf",
                file_type=".pdf",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0,
                page_count=page_count
            ),
            tables=tables,
            images=images
        )
    
    def _process_docx(self, file_path: Union[str, Path, bytes], 
                     filename: Optional[str] = None) -> ProcessedDocument:
        """Extract text from DOCX files"""
        
        try:
            if isinstance(file_path, bytes):
                doc = Document(io.BytesIO(file_path))
            else:
                doc = Document(file_path)
            
            content_parts = []
            tables = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    if row_idx == 0:
                        headers = row_data
                    else:
                        table_data.append(row_data)
                
                tables.append({
                    'table_index': table_idx,
                    'headers': headers,
                    'data': table_data
                })
            
            content = "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            content = f"Error processing DOCX: {str(e)}"
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "document.docx",
                file_type=".docx",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            ),
            tables=tables
        )
    
    def _process_doc(self, file_path: Union[str, Path, bytes], 
                    filename: Optional[str] = None) -> ProcessedDocument:
        """Process legacy DOC files (limited support)"""
        # Note: Full DOC support requires additional libraries like python-docx2txt
        try:
            import docx2txt
            if isinstance(file_path, bytes):
                # Save temporarily for docx2txt
                temp_path = f"/tmp/{filename or 'temp.doc'}"
                with open(temp_path, 'wb') as f:
                    f.write(file_path)
                content = docx2txt.process(temp_path)
                os.unlink(temp_path)
            else:
                content = docx2txt.process(str(file_path))
        except ImportError:
            content = "DOC file processing requires python-docx2txt library"
        except Exception as e:
            content = f"Error processing DOC: {str(e)}"
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "document.doc",
                file_type=".doc",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            )
        )
    
    def _process_csv(self, file_path: Union[str, Path, bytes], 
                    filename: Optional[str] = None) -> ProcessedDocument:
        """Process CSV files"""
        
        try:
            if isinstance(file_path, bytes):
                content_str = file_path.decode('utf-8')
                df = pd.read_csv(io.StringIO(content_str))
            else:
                df = pd.read_csv(file_path)
            
            # Generate summary
            summary_parts = [
                f"CSV Summary:",
                f"Rows: {len(df)}",
                f"Columns: {len(df.columns)}",
                f"Column Names: {', '.join(df.columns.tolist())}",
                "",
                "Data Types:",
            ]
            
            for col, dtype in df.dtypes.items():
                summary_parts.append(f"  {col}: {dtype}")
            
            summary_parts.extend([
                "",
                "Sample Data (first 5 rows):",
                df.head().to_string()
            ])
            
            content = "\n".join(summary_parts)
            
            # Structure the data
            structured_data = {
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.to_dict(),
                'shape': df.shape,
                'sample_data': df.head(5).to_dict('records'),
                'statistics': df.describe().to_dict() if len(df.select_dtypes(include='number').columns) > 0 else {}
            }
            
        except Exception as e:
            logger.error(f"CSV processing error: {str(e)}")
            content = f"Error processing CSV: {str(e)}"
            structured_data = None
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "data.csv",
                file_type=".csv",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            ),
            structured_data=structured_data
        )
    
    def _process_excel(self, file_path: Union[str, Path, bytes], 
                      filename: Optional[str] = None) -> ProcessedDocument:
        """Process Excel files (XLSX/XLS)"""
        
        try:
            if isinstance(file_path, bytes):
                excel_file = io.BytesIO(file_path)
            else:
                excel_file = file_path
            
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            content_parts = [f"Excel File Summary:"]
            content_parts.append(f"Number of sheets: {len(excel_data)}")
            content_parts.append("")
            
            structured_data = {'sheets': {}}
            
            for sheet_name, df in excel_data.items():
                content_parts.extend([
                    f"Sheet: {sheet_name}",
                    f"  Rows: {len(df)}",
                    f"  Columns: {len(df.columns)}",
                    f"  Column Names: {', '.join(df.columns.tolist())}",
                    "",
                    "Sample Data:",
                    df.head(3).to_string(),
                    ""
                ])
                
                structured_data['sheets'][sheet_name] = {
                    'columns': df.columns.tolist(),
                    'shape': df.shape,
                    'sample_data': df.head(5).to_dict('records')
                }
            
            content = "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"Excel processing error: {str(e)}")
            content = f"Error processing Excel: {str(e)}"
            structured_data = None
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "workbook.xlsx",
                file_type=".xlsx",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            ),
            structured_data=structured_data
        )
    
    def _process_text(self, file_path: Union[str, Path, bytes], 
                     filename: Optional[str] = None) -> ProcessedDocument:
        """Process plain text files"""
        
        try:
            if isinstance(file_path, bytes):
                content = file_path.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    if isinstance(file_path, bytes):
                        content = file_path.decode(encoding)
                    else:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                content = "Error: Could not decode text file"
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "text.txt",
                file_type=".txt",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            )
        )
    
    def _process_json(self, file_path: Union[str, Path, bytes], 
                     filename: Optional[str] = None) -> ProcessedDocument:
        """Process JSON files"""
        
        try:
            if isinstance(file_path, bytes):
                data = json.loads(file_path.decode('utf-8'))
            else:
                with open(file_path, 'r') as f:
                    data = json.load(f)
            
            # Generate readable summary
            content_parts = [
                "JSON Structure Analysis:",
                f"Type: {type(data).__name__}",
            ]
            
            if isinstance(data, dict):
                content_parts.append(f"Keys: {len(data)} - {', '.join(list(data.keys())[:10])}")
            elif isinstance(data, list):
                content_parts.append(f"Items: {len(data)}")
                if data and isinstance(data[0], dict):
                    content_parts.append(f"Sample keys: {', '.join(list(data[0].keys())[:5])}")
            
            content_parts.extend([
                "",
                "Pretty-printed JSON:",
                json.dumps(data, indent=2)[:2000] + ("..." if len(str(data)) > 2000 else "")
            ])
            
            content = "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"JSON processing error: {str(e)}")
            content = f"Error processing JSON: {str(e)}"
            data = None
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "data.json",
                file_type=".json",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            ),
            structured_data=data
        )
    
    def _process_xml(self, file_path: Union[str, Path, bytes], 
                    filename: Optional[str] = None) -> ProcessedDocument:
        """Process XML files"""
        
        try:
            if isinstance(file_path, bytes):
                xml_content = file_path.decode('utf-8')
            else:
                with open(file_path, 'r') as f:
                    xml_content = f.read()
            
            soup = BeautifulSoup(xml_content, 'xml')
            
            # Extract text content
            text_content = soup.get_text()
            
            # Structure analysis
            content_parts = [
                "XML Structure Analysis:",
                f"Root element: {soup.name if soup.name else 'Unknown'}",
                f"Total elements: {len(soup.find_all())}",
                "",
                "Text content:",
                text_content[:1000] + ("..." if len(text_content) > 1000 else "")
            ]
            
            content = "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"XML processing error: {str(e)}")
            content = f"Error processing XML: {str(e)}"
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "data.xml",
                file_type=".xml",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            )
        )
    
    def _process_html(self, file_path: Union[str, Path, bytes], 
                     filename: Optional[str] = None) -> ProcessedDocument:
        """Process HTML files"""
        
        try:
            if isinstance(file_path, bytes):
                html_content = file_path.decode('utf-8')
            else:
                with open(file_path, 'r') as f:
                    html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text()
            
            # Extract structured data
            structured_data = {
                'title': soup.title.string if soup.title else None,
                'headings': [h.get_text() for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])],
                'links': [{'text': a.get_text(), 'href': a.get('href')} for a in soup.find_all('a', href=True)[:10]],
                'images': [img.get('src') for img in soup.find_all('img', src=True)[:5]]
            }
            
            content_parts = [
                "HTML Content Analysis:",
                f"Title: {structured_data['title'] or 'No title'}",
                f"Headings found: {len(structured_data['headings'])}",
                f"Links found: {len(structured_data['links'])}",
                f"Images found: {len(structured_data['images'])}",
                "",
                "Text content:",
                text_content[:1500] + ("..." if len(text_content) > 1500 else "")
            ]
            
            content = "\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"HTML processing error: {str(e)}")
            content = f"Error processing HTML: {str(e)}"
            structured_data = None
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=filename or "page.html",
                file_type=".html",
                size_bytes=len(file_path) if isinstance(file_path, bytes) else 0
            ),
            structured_data=structured_data
        )
    
    def _process_url_direct(self, url: str) -> ProcessedDocument:
        """Process URL content directly"""
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            
            if 'html' in content_type:
                return self._process_html(response.content.encode('utf-8'), filename=url)
            elif 'json' in content_type:
                return self._process_json(response.content, filename=url)
            else:
                return self._process_text(response.content, filename=url)
                
        except Exception as e:
            logger.error(f"URL processing error: {str(e)}")
            return ProcessedDocument(
                content=f"Error processing URL {url}: {str(e)}",
                metadata=DocumentMetadata(
                    filename=url,
                    file_type=".url",
                    size_bytes=0
                )
            )
    
    def _process_url_serpapi(self, url: str) -> ProcessedDocument:
        """Process URL using SerpAPI for enhanced web scraping"""
        
        if not self.serpapi_key:
            return self._process_url_direct(url)
        
        try:
            search = GoogleSearch({
                "q": f"site:{url}",
                "api_key": self.serpapi_key
            })
            results = search.get_dict()
            
            # Extract organic results
            organic_results = results.get("organic_results", [])
            
            content_parts = [
                f"Web Content Analysis for: {url}",
                f"Results found: {len(organic_results)}",
                ""
            ]
            
            for result in organic_results[:5]:
                content_parts.extend([
                    f"Title: {result.get('title', 'No title')}",
                    f"Snippet: {result.get('snippet', 'No snippet')}",
                    f"Link: {result.get('link', 'No link')}",
                    ""
                ])
            
            content = "\n".join(content_parts)
            
            structured_data = {
                'results_count': len(organic_results),
                'results': organic_results
            }
            
        except Exception as e:
            logger.error(f"SerpAPI processing error: {str(e)}")
            return self._process_url_direct(url)
        
        return ProcessedDocument(
            content=content,
            metadata=DocumentMetadata(
                filename=url,
                file_type=".serpapi",
                size_bytes=len(content.encode('utf-8'))
            ),
            structured_data=structured_data
        )

# Integration with Sakthi Language
class SakthiDocumentProcessor:
    """Integration layer between document processing and Sakthi language"""
    
    def __init__(self, sakthi_engine, serpapi_key: Optional[str] = None):
        self.sakthi_engine = sakthi_engine
        self.doc_processor = DocumentProcessor(serpapi_key)
    
    def process_document_with_intent(self, file_path: Union[str, Path, io.BytesIO], 
                                   intent_text: str, 
                                   filename: Optional[str] = None) -> Dict[str, Any]:
        """Process document and apply Sakthi intent analysis"""
        
        # Process document
        processed_doc = self.doc_processor.process_file(file_path, filename)
        
        # Combine document content with intent
        combined_input = f"{intent_text}\n\nDocument Content:\n{processed_doc.content}"
        
        # Process with Sakthi
        sakthi_result = self.sakthi_engine.process(combined_input)
        
        return {
            'document': processed_doc,
            'sakthi_analysis': sakthi_result,
            'processing_timestamp': datetime.now().isoformat()
        }
    
    def batch_process_documents(self, file_paths: List[Union[str, Path]], 
                              intent_text: str) -> List[Dict[str, Any]]:
        """Process multiple documents with the same intent"""
        
        results = []
        for file_path in file_paths:
            try:
                result = self.process_document_with_intent(file_path, intent_text)
                results.append(result)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                results.append({
                    'error': str(e),
                    'file_path': str(file_path)
                })
        
        return results

# Example usage
if __name__ == "__main__":
    # Initialize processor
    processor = DocumentProcessor()
    
    # Test different file types
    test_files = [
        "sample.pdf",
        "data.csv", 
        "report.docx",
        "config.json"
    ]
    
    print("=== Document Processing Demo ===\n")
    
    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"Processing: {file_path}")
            try:
                result = processor.process_file(file_path)
                print(f"Content length: {len(result.content)}")
                print(f"Tables found: {len(result.tables) if result.tables else 0}")
                print(f"Confidence: {result.confidence}")
                print("-" * 50)
            except Exception as e:
                print(f"Error: {str(e)}")
        else:
            print(f"File not found: {file_path}")
    
    # Test URL processing
    test_url = "https://example.com"
    print(f"\nProcessing URL: {test_url}")
    try:
        url_result = processor.process_url(test_url)
        print(f"URL content length: {len(url_result.content)}")
    except Exception as e:
        print(f"URL processing error: {str(e)}")
    
    print("\n=== Processing Complete ===")